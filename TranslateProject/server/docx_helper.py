from docx import Document
from translator import translate_text

def process_docx_translation(input_path, output_path, target_lang):
    """
    Mở file docx, dịch nội dung bên trong và lưu lại với định dạng gốc.
    """
    doc = Document(input_path)

    # 1. Dịch các đoạn văn bản (Paragraphs)
    for p in doc.paragraphs:
        if p.text.strip():
            # Lưu ý: Việc gán p.text sẽ giữ nguyên style chung của paragraph 
            # nhưng có thể làm mất định dạng riêng lẻ (như 1 chữ in đậm giữa câu).
            # Để giữ style chi tiết hơn, cần dịch theo từng 'run', nhưng sẽ làm context dịch bị chia nhỏ.
            p.text = translate_text(p.text, target_lang)

    # 2. Dịch văn bản trong bảng (Tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = translate_text(cell.text, target_lang)

    # Lưu lại file với định dạng docx
    doc.save(output_path)